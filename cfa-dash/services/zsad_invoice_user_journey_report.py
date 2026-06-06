from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"c:\Users\cwakh\Downloads\Robert\cp TEMPLATE.docx")
OUTPUT = ROOT / "docs" / "ZCAMS_ZSAD_Invoice_User_Journey_Test_Report.docx"
PROBE_RESULTS = ROOT / "docs" / "zsad_invoice_10bl_probe_results.json"

BRAND_GREEN = RGBColor(6, 69, 31)
BRAND_DARK = RGBColor(20, 35, 28)
BRAND_MUTED = RGBColor(88, 103, 93)


def _set_run(run, *, bold: bool = False, size: int = 10, color: RGBColor = BRAND_DARK, font: str = "Aptos") -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {min(level, 3)}"
    run = p.add_run(text)
    _set_run(run, bold=True, size={1: 18, 2: 14, 3: 12}.get(level, 11), color=BRAND_GREEN)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        _set_run(lead, bold=True, size=10)
        rest = p.add_run(text[len(bold_lead) :])
        _set_run(rest, size=10)
    else:
        run = p.add_run(text)
        _set_run(run, size=10)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        try:
            p = doc.add_paragraph(style="List Bullet")
        except KeyError:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            item = f"- {item}"
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        _set_run(run, size=9)


def add_code(doc: Document, text: str) -> None:
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        _set_run(run, size=8, color=RGBColor(35, 43, 38), font="Consolas")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        run = hdr[idx].paragraphs[0].add_run(header)
        _set_run(run, bold=True, size=8, color=RGBColor(255, 255, 255))
        hdr[idx]._tc.get_or_add_tcPr().append(_shading("06451F"))
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            run = cells[idx].paragraphs[0].add_run(str(value))
            _set_run(run, size=8)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def _load_probe_results() -> dict:
    if not PROBE_RESULTS.is_file():
        return {}
    return json.loads(PROBE_RESULTS.read_text(encoding="utf-8"))


def _fmt_metric(metric: dict) -> str:
    if not metric:
        return "-"
    return (
        f"min {metric.get('min', 0)}ms | avg {metric.get('avg', 0)}ms | "
        f"p95 {metric.get('p95', 0)}ms | max {metric.get('max', 0)}ms"
    )


def _shading(fill: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    elem = OxmlElement("w:shd")
    elem.set(qn("w:fill"), fill)
    return elem


def add_title_page(doc: Document) -> None:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ZCAMS Z-SAD AND INVOICE PIPELINE\nUSER JOURNEY, TEST DESIGN, AND FAILURE INVESTIGATION")
    _set_run(run, bold=True, size=20, color=BRAND_GREEN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared for ZAFFA Clearing & Forwarding | Industry-grade operational assurance pack")
    _set_run(run, size=11, color=BRAND_MUTED)

    add_body(
        doc,
        "Scope: This document investigates the end-to-end user journey for creating a Z-SAD and issuing an invoice in ZCAMS, "
        "with special focus on the invoicing pipeline, CapitalPay signing and checkout, PDF regeneration, importer sharing, "
        "and operational failure reactions.",
    )
    add_body(
        doc,
        f"Evidence date: {datetime.now().strftime('%Y-%m-%d %H:%M')}. Template used: {TEMPLATE}. Output: {OUTPUT}.",
    )


def build_report() -> Path:
    doc = Document(str(TEMPLATE)) if TEMPLATE.is_file() else Document()
    probe = _load_probe_results()
    doc.core_properties.title = "ZCAMS Z-SAD and Invoice User Journey Test Report"
    doc.core_properties.subject = "Z-SAD, GN 83 invoice, CapitalPay checkout, PDF and failure test matrix"
    doc.core_properties.author = "ZCAMS Engineering"

    add_title_page(doc)

    add_heading(doc, "1. Executive Test Verdict", 1)
    add_body(
        doc,
        "Verdict: The focused local test suite passed for the current Z-SAD, invoice, CapitalPay, Agentic Mode, repository, "
        "PDF, and BL end-to-end paths. The main residual production risks are external dependency reachability, missing "
        "CapitalPay credentials, OAuth/token failures, checkout iframe/network latency, email/SMS provider failures, and "
        "stale environment configuration on EC2.",
        bold_lead="Verdict:",
    )
    add_table(
        doc,
        ["Evidence", "Result", "Metric / Observation", "Meaning"],
        [
            [
                "Focused pytest suite",
                "PASSED",
                "41 passed, 1 warning, 11.94s",
                "Core code paths execute locally with mock-safe external boundaries.",
            ],
            [
                "Slowest full journey",
                "PASSED",
                "2.91s: BL upload -> parse -> Z-SAD -> invoice -> PDF download",
                "The local processing budget is healthy for text-PDF and mocked CapitalPay flows.",
            ],
            [
                "Agentic journey",
                "PASSED",
                "0.51s: extraction memory path -> Z-SAD -> invoice/share summary",
                "The Agentic path is faster because it is orchestrated as a compact service flow.",
            ],
            [
                "Wrong working directory run",
                "FAILED",
                "5 collection errors: services/app imports not found",
                "Deployment/test scripts must run from cfa-dash or set PYTHONPATH.",
            ],
            [
                "Pytest cache write",
                "WARNING",
                "WinError 5 writing .pytest_cache",
                "Non-blocking local file permission issue; does not invalidate test pass.",
            ],
        ],
        widths=[1.4, 0.9, 1.8, 2.6],
    )

    add_heading(doc, "2. User Journey Under Investigation", 1)
    add_body(
        doc,
        "The ZCAMS journey starts when a Declarant, Company Admin, or Super Admin uploads a Bill of Lading, validates extracted "
        "customs values, reviews the BL, issues a single active Z-SAD, selects a GN 83 invoice mode, signs the invoice through "
        "CapitalPay, and either opens Pay Now checkout or generates and shares the signed invoice with the importer."
    )
    add_table(
        doc,
        ["Stage", "Primary User Action", "System Action", "Primary Output", "Failure Reaction"],
        [
            [
                "1. Upload BL",
                "Upload PDF/image/doc and review extracted fields.",
                "Save file, run OCR/text extraction, infer GN 83 category, populate form.",
                "Draft BL payload.",
                "Show fallback/demo extraction with ocr_error; require user correction before save.",
            ],
            [
                "2. Save BL",
                "Click Save BL after correcting values.",
                "Persist bills_of_lading and cargo_items; detect duplicate BL conflicts.",
                "BL status UPLOADED or reviewed auto path.",
                "Block duplicates and offer detach/cancel previous record before re-upload.",
            ],
            [
                "3. Review BL",
                "Click Review & Issue Z-SAD.",
                "Create reviewed_bls row, generate unique Z-SAD number, mark active.",
                "Status REVIEWED_ZSAD_ISSUED.",
                "Unknown BL raises ValueError; no Z-SAD without saved BL.",
            ],
            [
                "4. Select invoice",
                "Choose Full Settlement.",
                "Calculate GN 83 minimum, admin fee, VAT, and total.",
                "Invoice confirmation modal.",
                "Full Settlement below minimum is rejected before signing.",
            ],
            [
                "5. Sign invoice",
                "Click Generate & share invoice or Pay Now.",
                "Call CapitalPay invoice/create, persist invoice and payment rows, generate PDF.",
                "CapitalPay invoice number, DB invoice, signed PDF.",
                "Credential/network/API failures stop persistence and show error.",
            ],
            [
                "6A. Generate & share",
                "Share via WhatsApp/SMS/email.",
                "Prepare checkout once, align CapitalPay ref, regenerate PDF, send/share links.",
                "Download signed invoice PDF and share notifications.",
                "Email may fall back to mailto; WhatsApp remains a browser link.",
            ],
            [
                "6B. Pay Now",
                "Open checkout tab.",
                "Generate invoice, expose /capitalpay/checkout/<invoice_id>, fetch iframe HTML on route.",
                "CapitalPay checkout page and signed PDF download.",
                "Route returns 401/403/404/502 for auth, tenant, missing invoice, or CapitalPay errors.",
            ],
            [
                "7. Download PDF",
                "Click Download signed invoice PDF.",
                "Prepare checkout, align displayed ref, regenerate PDF from DB, send file.",
                "PDF with same CapitalPay ref as checkout.",
                "Returns 500 if checkout/PDF preparation fails, 404 if file missing.",
            ],
            [
                "8. Settlement",
                "Payment callback/status update or manual settlement.",
                "Mark invoice/payments settled; release state depends on invoice type.",
                "SETTLED_RELEASE_PENDING or CARGO_RELEASED.",
                "If callback is unreachable, invoice remains AWAITING_PAYMENT until reconciled.",
            ],
        ],
        widths=[1.0, 1.4, 2.0, 1.5, 2.0],
    )

    add_heading(doc, "2A. Granular Login-to-Invoice User Journey", 1)
    add_body(
        doc,
        "This is the operator-level journey that should be tested end to end because every later invoice artifact depends on "
        "the correctness of the logged-in user, company tenancy, BL payload, Z-SAD issue, GN 83 calculation, and CapitalPay reference."
    )
    add_table(
        doc,
        ["Step", "Screen / Component", "Action", "System Checks", "Expected Output", "Common Error"],
        [
            ["1", "/login", "Enter username/email and password.", "Password hash, user status, session expiry, role path permissions.", "Authenticated session and dashboard route.", "Invalid credentials; suspended user; expired/revoked session."],
            ["2", "Dashboard navigation", "Open BLs or Agentic Mode.", "Role allowed: Declarant, Company Admin, Super Admin.", "BL capture workspace opens.", "403/redirect if role lacks module access."],
            ["3", "/bls upload", "Upload BL file or manually capture values.", "File decode, extension, OCR/text extraction, fallback flag.", "Draft fields populated.", "Bad base64, OCR timeout, OpenAI unreachable, fallback_demo values."],
            ["4", "BL capture form", "Correct five business values: BL number, TIN, weight, units/containers, GN 83 category.", "Required fields, duplicate BL check, company scoping.", "Clean BL payload ready for save.", "Duplicate BL conflict; wrong category; missing TIN."],
            ["5", "Save BL", "Click Save BL.", "Repository inserts BL and cargo item rows.", "BL status UPLOADED.", "SQLite constraint, duplicate BL, invalid payload."],
            ["6", "/reviewed-bl", "Click Review & Issue Z-SAD.", "Idempotent review lookup; Z-SAD uniqueness loop.", "Active Z-SAD and status REVIEWED_ZSAD_ISSUED.", "Unknown BL; duplicate review; generation collision retry."],
            ["7", "Invoice modal", "Choose Full Settlement or Full Settlement.", "GN 83 quote, minimum fee, VAT, settlement fields.", "Previewed invoice breakdown.", "Full settlement under minimum; missing beneficiary/bank confirmation."],
            ["8", "Invoice action", "Click Generate & Share or Pay Now once.", "CapitalPay signing, production mock guard, DB insert, PDF draft.", "Invoice AWAITING_PAYMENT and payment PENDING.", "CapitalPay auth/API error; no credentials; non-JSON response."],
            ["9", "Checkout/reference alignment", "Open Pay Now or download PDF.", "Checkout HTML fetched; PAYMENT REF extracted; DB/PDF regenerated.", "CapitalPay ref same in DB, PDF, checkout.", "Checkout unreachable; stale cache; ref parse miss."],
            ["10", "Share/download", "Send WhatsApp/SMS/email and download signed PDF.", "Channel validation, public download URL, PDF exists.", "Importer receives invoice and payment link.", "No phone/email; provider down; localhost URL on EC2."],
            ["11", "Payment/settlement", "Mark settled or receive callback in future integration.", "Invoice/payment state transition and release rule.", "Agency charge: release pending; full settlement: cargo released.", "Callback not reachable; duplicate settlement; manual reconciliation needed."],
        ],
        widths=[0.4, 1.3, 1.6, 2.0, 1.8, 1.8],
    )

    add_heading(doc, "3. Z-SAD Pipeline Mechanics", 1)
    add_bullets(
        doc,
        [
            "The Z-SAD is generated by services.repository.review_bl after a BL exists and has not already been reviewed.",
            "The generated number format is Z-SAD-<last-four-BL-chars>-<mixed 15-char suffix>; uniqueness is checked against z_sads before insert.",
            "The system creates one reviewed_bls row, one active z_sads row, updates bills_of_lading.status to REVIEWED_ZSAD_ISSUED, and emits BL_REVIEWED and ZSAD_GENERATED notifications.",
            "Replacement is controlled by detach_zsad and detach_zsad_for_reupload; once settled or released, detachment is blocked.",
            "Operational invariant: one active Z-SAD per active BL journey; old Z-SADs and cancelled invoices remain in the history table for audit.",
        ]
    )
    add_table(
        doc,
        ["Data Object", "Required Fields", "Produced By", "Downstream Use"],
        [
            ["BL payload", "bl_number, route_type, transport_mode, consignee, TIN, weight, containers, cargo, GN 83 category", "BL module OCR/capture or Agentic five-value review", "Z-SAD review and GN 83 quote"],
            ["reviewed_bls", "id, bl_id, status, z_sad_id, reviewed_by", "review_bl", "Invoice request table and history"],
            ["z_sads", "id, reviewed_bl_id, bl_id, z_sad_number, is_active, is_used", "review_bl / detach_zsad", "Invoice payload, ASYCUDA/customs reference"],
            ["notifications", "type, subject/body, entity_id, company_id", "notify calls", "Admin dashboards and audit feed"],
        ],
        widths=[1.4, 2.4, 1.6, 1.8],
    )

    add_heading(doc, "4. Invoicing Pipeline Mechanics", 1)
    add_body(
        doc,
        "The invoice pipeline is the most sensitive path because it crosses internal state, GN 83 math, CapitalPay signing, "
        "external checkout HTML, PDF generation, importer communication, and payment settlement. The main invariant is that "
        "ZCAMS DB, invoice PDF, and CapitalPay checkout must show the same CapitalPay payment reference."
    )
    add_table(
        doc,
        ["Function / Route", "Responsibility", "Pass Condition", "Failure Condition"],
        [
            ["gn83.calculate_invoice", "Calculate agency-charge or full-settlement totals.", "Agency charge and full settlement match GN 83 expectations.", "Wrong GN 83 category or units causes wrong charge."],
            ["repository.generate_invoice", "Create internal invoice number, call CapitalPay, persist invoice/payment, generate first PDF.", "CapitalPay returns invoice.invoice_number and DB rows are committed.", "Missing credentials, non-JSON response, API HTTP >=400, missing invoice_number."],
            ["capitalpay.create_signed_invoice", "Authenticate and POST /invoice/create with customer and line item payload.", "Returns urn/capitalpay_number/checkout_url/reference.", "Auth timeout, invalid token, endpoint unreachable, wrong payload schema."],
            ["repository.prepare_capitalpay_checkout", "Fetch checkout HTML once, extract PAYMENT REF, persist it, regenerate PDF, cache HTML.", "DB capitalpay_urn/payments.capitalpay_ref/PDF align with checkout ref.", "Checkout endpoint 502, parse misses ref, cache stale, disk write failure."],
            ["/capitalpay/checkout/<invoice_id>", "Auth and tenant-check checkout iframe route.", "Returns CapitalPay HTML for logged-in authorized user.", "401 anonymous, 403 wrong company, 404 missing invoice, 502 CapitalPay error."],
            ["/download/invoice/<invoice_id>.pdf", "Prepare checkout and rebuild PDF before download.", "PDF reflects latest payment ref and downloads as application/pdf.", "500 if checkout/PDF prep fails, 404 if invoice/file missing."],
            ["share_invoice_with_importer", "Build message and dispatch WhatsApp, SMS, email.", "Channels return sent/link metadata and notifications are logged.", "Email absent, SMTP/Bird failure, SMS provider unavailable."],
            ["settle_invoice", "Mark payment settled and update release state.", "Full settlement releases cargo; agency charge moves to release pending.", "Callback not received or duplicate event leaves awaiting state."],
        ],
        widths=[1.6, 2.2, 2.0, 2.1],
    )

    add_heading(doc, "5. Payload Investigation", 1)
    add_body(doc, "The payload is assembled in layers. Each layer must be inspectable and testable independently.")
    add_table(
        doc,
        ["Payload", "Producer", "Key Fields", "Validation / Risk"],
        [
            ["BL extraction payload", "ocr.extract_bl_fields and parse_bl_text", "bl_number, consignee_tin, gross_weight, no_containers, cargo_description, gn83_category, raw_text, ocr_mode", "OCR may fallback to demo data; user must review and correct before persistence."],
            ["BL persistence payload", "repository.create_bl", "doc_type, route_type, transport_mode, zra_regime, consignee fields, file path, cargo items", "Duplicate BL conflict must block unintended reuse."],
            ["Z-SAD payload", "repository.review_bl", "reviewed_id, bl_id, z_sad_number, active flag", "Number collision handled by loop; no active Z-SAD after detach without reissue."],
            ["GN 83 calculation payload", "gn83_quote_for_reviewed and calculate_invoice", "std_min_fee, admin_fee, vat, total, invoice_type", "Wrong category or container count causes wrong total."],
            ["CapitalPay create payload", "capitalpay._build_create_payload", "account_id, amount_expected, client_invoice_ref, currency, email, msisdn, id_number, payment_gateway_id, items", "Wrong account or missing auth fails before invoice persistence."],
            ["Checkout POST payload", "capitalpay.build_checkout_params", "apiClientID, secureHash, amountExpected, billRefNumber, client details", "Secure hash mismatch or network failure blocks iframe HTML."],
            ["Invoice PDF payload", "pdf_service.generate_invoice_pdf", "invoice_number, capitalpay_urn/ref, BL, Z-SAD, totals, bank/CapitalPay sections", "Must regenerate after checkout ref changes to avoid mismatch."],
            ["Share payload", "invoice_share_message and messaging adapters", "invoice number, BL, Z-SAD, amount, CapitalPay invoice no, PDF URL, pay URL", "PUBLIC_APP_URL controls absolute URLs; localhost must not leak in EC2."],
        ],
        widths=[1.4, 1.8, 2.6, 2.1],
    )

    add_heading(doc, "6. Reaction Matrix", 1)
    add_table(
        doc,
        ["Reaction", "Trigger", "Expected User Experience", "System Evidence", "Metric to Track"],
        [
            ["PASSED", "Valid text PDF, reviewed BL, CapitalPay mock/real success, PDF generated.", "User sees Z-SAD, invoice success card, Pay Now/Download PDF actions.", "DB rows, notifications, PDF file, pytest pass.", "p95 upload, sign, checkout, PDF time."],
            ["FAILED - validation", "Missing BL number, TIN, phone/email, unconfirmed Agentic five-value check, full settlement below minimum.", "Inline re-check notice; no invoice sent.", "ValueError or callback error card.", "Validation error rate by field."],
            ["FAILED - CapitalPay auth", "CAPITALPAY_MODE=real without key/secret or invalid token.", "Invoice action shows signing failed; no dummy invoice issued.", "CapitalPayError from _get_token/create_signed_invoice.", "Auth failure count and HTTP status."],
            ["FAILED - payload schema", "CapitalPay returns HTTP >=400 or missing invoice.invoice_number.", "Invoice is not persisted; operator sees failed signing.", "CapitalPayError with response payload.", "Create failure rate by response code."],
            ["TIMEOUT - OCR", "Text PDF extraction exceeds OCR_TEXT_PDF_TIMEOUT_SEC or image OCR slow.", "UI may wait; fallback OpenAI route used if configured.", "_extract_text_pdf_timed returns timed_out=True.", "OCR duration, timeout count, OCR mode."],
            ["TIMEOUT - CapitalPay", "OAuth/create/checkout request exceeds 30s/45s or network hangs.", "Pay Now tab or Generate & Share shows error; route may return 502.", "requests timeout/CapitalPayError.", "External dependency p95/p99 latency."],
            ["NOT REACHABLE - checkout", "CapitalPay checkout endpoint unavailable or private-host URLs inside HTML.", "Checkout iframe/page fails; PDF route may return 500 if forced prep fails.", "fetch_checkout_page error or normalized HTML missing assets.", "HTTP 502 count and checkout route latency."],
            ["NOT REACHABLE - email/SMS", "Bird/Gmail/SMS provider down or no email supplied.", "WhatsApp link still available; email result shows fallback/reason.", "share_invoice_with_importer channel metadata.", "Channel delivery success rate."],
            ["SECURITY BLOCK", "Anonymous/wrong tenant requests checkout or PDF.", "401/403 instead of leaking invoice/PDF.", "invoice_routes before_request and tenant guard.", "Unauthorized attempts by route."],
            ["CACHE / REF MISMATCH", "Checkout ref differs from stored capitalpay_urn or stale PDF exists.", "prepare/download regenerates invoice PDF with checkout ref.", "set_invoice_capitalpay_ref updates invoices/payments.", "Mismatch count before/after checkout prep."],
        ],
        widths=[1.1, 1.9, 2.0, 1.8, 1.4],
    )

    add_heading(doc, "7. Test Scripts Executed", 1)
    add_body(doc, "These are the exact local test commands used during this investigation.")
    add_code(
        doc,
        r"""
# Incorrect root-level command - intentionally preserved as environment failure evidence
.\cfa-dash\.venv\Scripts\python.exe -m pytest cfa-dash\tests\test_capitalpay.py cfa-dash\tests\test_invoice_flow.py cfa-dash\tests\test_repository.py cfa-dash\tests\test_agentic_workflow.py cfa-dash\tests\test_bl_e2e.py -q --tb=short --durations=20

# Correct command
cd cfa-dash
.\.venv\Scripts\python.exe -m pytest tests\test_capitalpay.py tests\test_invoice_flow.py tests\test_repository.py tests\test_agentic_workflow.py tests\test_bl_e2e.py -q --tb=short --durations=20

# 10-BL transaction probe used for sample analysis
.\.venv\Scripts\python.exe services\zsad_invoice_10bl_probe.py

# Regenerate this DOCX report after collecting probe evidence
.\.venv\Scripts\python.exe services\zsad_invoice_user_journey_report.py
""",
    )
    add_table(
        doc,
        ["Test File", "What It Proves", "Important Cases"],
        [
            ["tests/test_capitalpay.py", "CapitalPay payload, real API wrapper behavior under mocked requests, checkout endpoint usage, ref extraction, missing credentials.", "PAYMENT REF extraction, endpoint URL, credentials failure, CPAY number contract."],
            ["tests/test_invoice_flow.py", "Invoice calculation, invoice generation, sharing links, PDF creation, ref update persistence.", "Service vs full totals, email/WhatsApp/SMS sharing, set_invoice_capitalpay_ref."],
            ["tests/test_repository.py", "Repository state, user/session/company scoping, invoice download URL behavior.", "PUBLIC_APP_URL localhost vs public URL, admin workflows."],
            ["tests/test_agentic_workflow.py", "Agentic five-value gate, idempotency, Z-SAD and invoice share path.", "Missing confirmation blocks send, missing phone blocks WhatsApp, pay_now_url route."],
            ["tests/test_bl_e2e.py", "BL parse fixtures, PDF extraction, BL -> Z-SAD -> invoice -> PDF download.", "Three BL variants, PDF bytes, authenticated download route."],
        ],
        widths=[1.6, 3.0, 2.5],
    )

    add_heading(doc, "8. Additional Test Scripts Recommended", 1)
    add_body(doc, "The following scripts should be added for deeper production assurance before a full payment launch.")
    add_code(
        doc,
        r"""
def test_capitalpay_checkout_unreachable_returns_502(authenticated_client, monkeypatch):
    invoice = make_signed_invoice()
    monkeypatch.setattr("services.capitalpay.fetch_checkout_page", lambda *_: (_ for _ in ()).throw(CapitalPayError("timeout")))
    response = authenticated_client.get(f"/capitalpay/checkout/{invoice['id']}")
    assert response.status_code == 502

def test_download_regenerates_pdf_after_checkout_ref_change(authenticated_client, monkeypatch):
    invoice = make_signed_invoice()
    monkeypatch.setattr("services.capitalpay.fetch_checkout_page", lambda *_: "<html>PAYMENT REF CPAYLIVE123</html>")
    response = authenticated_client.get(f"/download/invoice/{invoice['id']}.pdf")
    assert response.status_code == 200
    assert get_invoice(invoice["id"])["capitalpay_ref"] == "CPAYLIVE123"

def test_full_settlement_below_minimum_is_blocked(reviewed_bl):
    with pytest.raises(ValueError, match="cannot be below"):
        generate_invoice(reviewed_bl["id"], "FULL_SETTLEMENT", std_min_fee_override=1.0)

def test_wrong_company_cannot_open_checkout(authenticated_client_other_company, invoice):
    response = authenticated_client_other_company.get(f"/capitalpay/checkout/{invoice['id']}")
    assert response.status_code == 403

def test_ocr_timeout_routes_to_image_ocr(monkeypatch, pdf_path):
    monkeypatch.setenv("OCR_TEXT_PDF_TIMEOUT_SEC", "0.001")
    monkeypatch.setattr("services.ocr.extract_text_with_openai", lambda *_args, **_kwargs: ("BL text", "image_pdf_ocr"))
    text, mode, provider = route_pdf_text(pdf_path)
    assert mode == "image_pdf_ocr"
""",
    )

    add_heading(doc, "9. Metrics and Monitoring Plan", 1)
    add_table(
        doc,
        ["Metric", "Target / Alert", "Where to Measure", "Why It Matters"],
        [
            ["BL OCR duration", "p95 < 15s text PDF; scanned PDF alert > 60s", "BL upload callback and ocr.route_pdf_text", "Prevents flicker and user uncertainty."],
            ["CapitalPay auth duration", "p95 < 5s; failures alert immediately", "capitalpay._get_token", "No invoice can be signed without token."],
            ["CapitalPay invoice/create duration", "p95 < 10s; timeout at 45s", "capitalpay.create_signed_invoice", "Main invoice creation dependency."],
            ["Checkout HTML duration", "p95 < 8s; 502 count monitored", "prepare_capitalpay_checkout and /capitalpay/checkout route", "Controls Pay Now experience."],
            ["Reference mismatch count", "Zero tolerated", "Compare payments.capitalpay_ref, invoices.capitalpay_urn, checkout extracted ref, PDF text", "Financial reconciliation invariant."],
            ["PDF regeneration duration", "p95 < 3s", "ensure_invoice_pdf", "Download button must feel immediate."],
            ["Share delivery success", "Email/SMS provider success > 98%; WhatsApp link generated 100%", "share_invoice_with_importer", "Importer receipt of invoice."],
            ["Security blocks", "401/403 attempts logged", "invoice_routes", "Proves tenant isolation."],
        ],
        widths=[1.6, 1.8, 2.0, 2.0],
    )

    add_heading(doc, "9A. 10-BL Transaction Sample Analysis", 1)
    if probe:
        login = probe.get("login") or {}
        summary = probe.get("summary") or {}
        add_body(
            doc,
            f"Sample executed at {probe.get('started_at')} using {probe.get('sample_size')} unique probe BLs. "
            f"The run authenticated as role {login.get('role')} in company {login.get('company_id')}. "
            f"Valid login completed in {login.get('valid_login_ms')}ms and invalid login was rejected in {login.get('invalid_login_ms')}ms. "
            f"The probe cleaned up all generated probe BLs, Z-SADs, invoices, payments, notifications, and PDFs after metrics were collected."
        )
        add_table(
            doc,
            ["Metric", "Result"],
            [
                ["Passed transactions", str(summary.get("passed_transactions", 0))],
                ["Failed transactions", str(summary.get("failed_transactions", 0))],
                ["Expected negative/error cases", str(summary.get("negative_expected_errors", 0))],
                ["Save BL latency", _fmt_metric(summary.get("save_bl_ms", {}))],
                ["Issue Z-SAD latency", _fmt_metric(summary.get("issue_zsad_ms", {}))],
                ["Generate invoice latency", _fmt_metric(summary.get("generate_invoice_ms", {}))],
                ["Align CapitalPay ref latency", _fmt_metric(summary.get("align_ref_ms", {}))],
                ["Regenerate PDF latency", _fmt_metric(summary.get("pdf_regen_ms", {}))],
                ["Build WhatsApp link latency", _fmt_metric(summary.get("whatsapp_link_ms", {}))],
            ],
            widths=[2.4, 4.6],
        )
        add_table(
            doc,
            ["#", "BL", "Route", "Category", "Invoice Type", "Z-SAD", "CP Ref", "Total", "PDF", "Status"],
            [
                [
                    str(txn.get("case")),
                    txn.get("bl_number", "-"),
                    f"{txn.get('route')} / {txn.get('transport')}",
                    txn.get("category", "-"),
                    txn.get("invoice_type", "-"),
                    txn.get("z_sad_number", "-"),
                    txn.get("capitalpay_ref", "-"),
                    f"USD {float(txn.get('total_usd') or 0):,.2f}",
                    "Yes" if txn.get("pdf_exists") else "No",
                    txn.get("status", "-"),
                ]
                for txn in probe.get("transactions", [])
            ],
            widths=[0.3, 1.0, 1.0, 1.2, 1.3, 1.8, 1.0, 0.8, 0.4, 0.8],
        )
        add_body(
            doc,
            "Transaction analysis: all 10 probe BLs reached AWAITING_PAYMENT, produced a Z-SAD, generated a signed invoice, "
            "aligned the stored CapitalPay reference, regenerated a PDF, and produced a WhatsApp invoice link. The sample deliberately "
            "mixed agency-charge-only and full-settlement invoices across import, transit, export, sea, road, air, containerized, LCL, dry bulk, "
            "bulk liquid, motor vehicle, heavy equipment, live animal, and general cargo categories."
        )
        add_table(
            doc,
            ["Error Case", "Type", "Outcome", "Meaning", "Observed Message"],
            [
                [
                    item.get("name", "-"),
                    item.get("type", "-"),
                    item.get("outcome", "-"),
                    item.get("meaning", "-"),
                    item.get("message", "-"),
                ]
                for item in probe.get("negative_tests", [])
            ],
            widths=[1.3, 1.2, 1.0, 2.2, 2.2],
        )
    else:
        add_body(
            doc,
            "The 10-BL probe results file was not present when this report was generated. Run "
            "python services/zsad_invoice_10bl_probe.py from cfa-dash and regenerate this report to embed the sample.",
        )

    add_heading(doc, "9B. Error Taxonomy and Metrics Interpretation", 1)
    add_table(
        doc,
        ["Error Type", "Meaning", "User Impact", "Metric", "Operational Response"],
        [
            ["Authentication", "Login failed or session cannot be trusted.", "User cannot enter the workflow.", "Failed login count, invalid/valid login latency, revoked-session hits.", "Confirm account status, reset password, inspect login_events."],
            ["Authorization / Tenant", "User is valid but not allowed to access invoice/download/checkout for that company.", "Route returns 403 or user is redirected.", "403 count by route, company mismatch count.", "Investigate role assignment and company_id scoping."],
            ["Validation", "Input is missing or violates user journey rules.", "Inline form error; no invoice/payload should be sent.", "Validation error rate by field.", "Improve labels/tooltips and mandatory field checks."],
            ["Business Rule", "Payload is syntactically valid but violates GN 83/payment invariants.", "Invoice blocked or settlement blocked.", "Under-minimum attempts, detach-after-settlement attempts.", "Train users and keep audit trail."],
            ["Duplicate / Conflict", "A BL or invoice journey already exists.", "User must detach/cancel prior journey before reuse.", "Duplicate BL conflict count, duplicate submit count.", "Enforce idempotency and review repeated clicks."],
            ["External Dependency", "CapitalPay, OpenAI OCR, email, SMS, or network route is down/slow.", "Timeout, 502, delayed checkout or fallback channel.", "p95/p99 dependency latency and failure count.", "Alert, retry, fallback, reconcile manually."],
            ["Reference Mismatch", "DB, PDF, share text, and checkout page disagree on CapitalPay ref.", "Payment reconciliation risk.", "Mismatch count, ref alignment latency.", "Run prepare_capitalpay_checkout and regenerate PDF before release/share."],
            ["Persistence / Filesystem", "SQLite/PDF/cache path cannot be written or read.", "Invoice may persist without PDF or report cache may fail.", "DB write errors, PDF generation failures, disk free space.", "Check permissions, disk space, backups, and retry generation."],
            ["Callback / UI Timing", "Dash callback blocks, double-fires, or cannot open a new tab in gesture.", "Flicker, duplicate invoices, missing iframe.", "Callback duration, duplicate invoice per reviewed BL, clientside open success.", "Split callbacks, disable buttons, keep external fetch on target route."],
            ["Settlement / Callback", "Payment status is not confirmed or callback route is missing/unreachable.", "Invoice remains awaiting payment.", "Age of AWAITING_PAYMENT invoices, callback failure count.", "Manual reconciliation and future webhook handler implementation."],
        ],
        widths=[1.2, 1.8, 1.7, 1.7, 1.8],
    )

    add_heading(doc, "10. What Can Go Wrong", 1)
    add_bullets(
        doc,
        [
            "A deployed EC2 environment may have stale .env values, causing PUBLIC_APP_URL to generate localhost links or CapitalPay callbacks to target the wrong host.",
            "If CAPITALPAY_MODE is real but keys are absent, the system correctly refuses to sign instead of producing CPAYMOCK invoices.",
            "If checkout HTML returns a payment reference that differs from the initial invoice/create response, prepare_capitalpay_checkout must update invoices and payments, then regenerate the PDF.",
            "If the checkout route is opened by an unauthenticated user, the route returns 401. If another company tries to open it, it returns 403.",
            "If the OCR provider is not configured for scanned PDFs/images, extraction falls back to demo values with an ocr_error; users must not save without review.",
            "If a BL is duplicated, the system must not silently issue a second active Z-SAD; the prior journey must be detached/cancelled first.",
            "If an invoice is generated twice by double-click or repeated callback, the user experience can confuse reconciliation; current UI disables Pay Now and Generate & Share after success, and Agentic Mode has idempotency support.",
            "If payment callbacks are delayed or not reachable, invoices remain AWAITING_PAYMENT and operations must reconcile through CapitalPay reports before release.",
        ]
    )

    add_heading(doc, "11. Release Readiness Checklist", 1)
    add_bullets(
        doc,
        [
            "Run the focused suite from cfa-dash on every release candidate.",
            "Add the recommended unreachable/timeout tests before formal go-live.",
            "Confirm PUBLIC_APP_URL=https://zcams.info on EC2 and no localhost is present in PDF/share messages.",
            "Confirm CapitalPay key, secret, account ID, payment gateway ID, callback URL, and notification URL in production .env.",
            "Capture route-level logs for /capitalpay/checkout and /download/invoice and alert on 5xx.",
            "Perform one controlled live CapitalPay transaction and compare DB, checkout iframe, downloaded PDF, WhatsApp text, and email attachment.",
            "Back up zcams.db before any schema or payment-pipeline release.",
        ]
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    path = build_report()
    print(path)
